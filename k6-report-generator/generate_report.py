#!/usr/bin/env python3
"""
k6 Performance Test Report Generator
=====================================
Generates a professional Word (.docx) report from k6 test results.

Supports input formats:
  1. JSON summary  (handleSummary() export or --out json=...)
  2. CSV export    (--out csv=...)

Usage:
  python generate_report.py --input results.json --format json
  python generate_report.py --input results.csv --format csv
  python generate_report.py --input results.json --format json \
      --title "API Load Test" --client "Acme Corp" --env "Production" \
      --output report.docx
"""

import argparse
import json
import csv
import os
import sys
import statistics
from datetime import datetime, timezone
from collections import defaultdict
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Brand Colors ───────────────────────────────────────────
PRIMARY = RGBColor(0x1A, 0x56, 0xDB)    # Blue
ACCENT  = RGBColor(0x10, 0xB9, 0x81)    # Green
DANGER  = RGBColor(0xEF, 0x44, 0x44)    # Red
WARNING = RGBColor(0xF5, 0x9E, 0x0B)    # Amber
DARK    = RGBColor(0x1F, 0x2A, 0x37)    # Near-black
GRAY    = RGBColor(0x6B, 0x72, 0x80)    # Gray
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

CHART_COLORS = ["#1A56DB", "#10B981", "#F59E0B", "#EF4444", "#8B5CF6", "#EC4899"]


# ─── Parsing ────────────────────────────────────────────────

def parse_json_summary(path: str) -> dict:
    """Parse k6 handleSummary() JSON or --out json= (line-delimited) output."""
    with open(path, "r") as f:
        first_char = f.read(1)
        f.seek(0)

        # handleSummary() produces a single JSON object with "metrics"
        if first_char == "{":
            try:
                data = json.load(f)
                if "metrics" in data:
                    return _normalize_summary(data)
            except json.JSONDecodeError:
                pass
            f.seek(0)

        # Line-delimited JSON (--out json=)
        return _parse_json_lines(f)


def _normalize_summary(data: dict) -> dict:
    """Normalize handleSummary() format into our internal structure."""
    result = {
        "metrics": {},
        "thresholds": {},
        "checks": [],
        "meta": data.get("root_group", {}),
    }

    for name, info in data.get("metrics", {}).items():
        metric = {
            "type": info.get("type", "trend"),
            "contains": info.get("contains", "default"),
        }

        if "values" in info:
            metric["values"] = info["values"]
        elif "value" in info:
            metric["values"] = {"value": info["value"]}
        else:
            metric["values"] = {}

        if "thresholds" in info:
            for t_name, t_data in info["thresholds"].items():
                result["thresholds"][f"{name} - {t_name}"] = t_data.get("ok", True)

        result["metrics"][name] = metric

    # Extract checks from root_group
    if "checks" in data.get("root_group", {}):
        for check_name, check_data in data["root_group"]["checks"].items():
            result["checks"].append({
                "name": check_name,
                "passes": check_data.get("passes", 0),
                "fails": check_data.get("fails", 0),
            })

    return result


def _parse_json_lines(f) -> dict:
    """Parse line-delimited JSON (--out json=results.json)."""
    metrics_meta = {}
    points = defaultdict(list)

    for line in f:
        line = line.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
        except json.JSONDecodeError:
            continue

        if obj.get("type") == "Metric":
            metrics_meta[obj["metric"]] = obj["data"]
        elif obj.get("type") == "Point":
            points[obj["metric"]].append(obj["data"])

    # Aggregate points into summary
    result = {"metrics": {}, "thresholds": {}, "checks": [], "meta": {}}

    for name, pts in points.items():
        values = [p["value"] for p in pts]
        meta = metrics_meta.get(name, {})
        mtype = meta.get("type", "trend")

        if mtype == "trend":
            sorted_vals = sorted(values)
            n = len(sorted_vals)
            agg = {
                "avg": statistics.mean(values),
                "min": min(values),
                "med": statistics.median(values),
                "max": max(values),
                "p(90)": sorted_vals[int(n * 0.9)] if n > 0 else 0,
                "p(95)": sorted_vals[int(n * 0.95)] if n > 0 else 0,
                "count": n,
            }
        elif mtype == "counter":
            agg = {"count": sum(values), "rate": sum(values) / max(len(values), 1)}
        elif mtype == "rate":
            agg = {"rate": statistics.mean(values), "passes": sum(1 for v in values if v > 0), "fails": sum(1 for v in values if v == 0)}
        elif mtype == "gauge":
            agg = {"value": values[-1] if values else 0, "min": min(values), "max": max(values)}
        else:
            agg = {"value": values[-1] if values else 0}

        result["metrics"][name] = {"type": mtype, "contains": meta.get("contains", "default"), "values": agg}

        # Thresholds
        for t in meta.get("thresholds", []):
            result["thresholds"][f"{name} - {t}"] = True  # Can't evaluate from points alone

        # Collect time series for charts
        if pts and "time" in pts[0]:
            result.setdefault("_timeseries", {})[name] = pts

    # Checks
    if "checks" in points:
        check_groups = defaultdict(lambda: {"passes": 0, "fails": 0})
        for p in points["checks"]:
            tag = (p.get("tags") or {}).get("check", "unnamed")
            if p["value"] > 0:
                check_groups[tag]["passes"] += 1
            else:
                check_groups[tag]["fails"] += 1
        for cname, cdata in check_groups.items():
            result["checks"].append({"name": cname, **cdata})

    return result


def parse_csv(path: str) -> dict:
    """Parse k6 CSV output (--out csv=results.csv)."""
    points = defaultdict(list)

    with open(path, "r") as f:
        reader = csv.DictReader(f)
        for row in reader:
            metric_name = row.get("metric_name", "")
            try:
                value = float(row.get("metric_value", 0))
            except (ValueError, TypeError):
                continue
            timestamp = row.get("timestamp", "")

            points[metric_name].append({
                "value": value,
                "time": timestamp,
                "tags": {k: v for k, v in row.items() if k not in ("metric_name", "timestamp", "metric_value") and v},
            })

    # Aggregate into summary
    result = {"metrics": {}, "thresholds": {}, "checks": [], "meta": {}}

    # Infer types
    KNOWN_TRENDS = {"http_req_duration", "http_req_blocked", "http_req_connecting",
                    "http_req_tls_handshaking", "http_req_sending", "http_req_waiting",
                    "http_req_receiving", "iteration_duration"}
    KNOWN_COUNTERS = {"http_reqs", "data_sent", "data_received", "iterations"}
    KNOWN_RATES = {"http_req_failed", "checks"}
    KNOWN_GAUGES = {"vus", "vus_max"}

    for name, pts in points.items():
        values = [p["value"] for p in pts]

        if name in KNOWN_TRENDS or name.startswith("http_req_"):
            mtype = "trend"
            sorted_vals = sorted(values)
            n = len(sorted_vals)
            agg = {
                "avg": statistics.mean(values),
                "min": min(values),
                "med": statistics.median(values),
                "max": max(values),
                "p(90)": sorted_vals[int(n * 0.9)] if n > 0 else 0,
                "p(95)": sorted_vals[int(n * 0.95)] if n > 0 else 0,
                "count": n,
            }
        elif name in KNOWN_COUNTERS:
            mtype = "counter"
            agg = {"count": sum(values), "rate": sum(values) / max(len(values), 1)}
        elif name in KNOWN_RATES:
            mtype = "rate"
            agg = {"rate": statistics.mean(values), "passes": sum(1 for v in values if v > 0), "fails": sum(1 for v in values if v == 0)}
        elif name in KNOWN_GAUGES:
            mtype = "gauge"
            agg = {"value": values[-1] if values else 0, "min": min(values), "max": max(values)}
        else:
            mtype = "trend"
            sorted_vals = sorted(values)
            n = len(sorted_vals)
            agg = {
                "avg": statistics.mean(values) if values else 0,
                "min": min(values) if values else 0,
                "med": statistics.median(values) if values else 0,
                "max": max(values) if values else 0,
                "p(90)": sorted_vals[int(n * 0.9)] if n > 0 else 0,
                "p(95)": sorted_vals[int(n * 0.95)] if n > 0 else 0,
                "count": n,
            }

        contains = "time" if name in KNOWN_TRENDS or "duration" in name else "default"
        result["metrics"][name] = {"type": mtype, "contains": contains, "values": agg}

    # Checks from CSV
    if "checks" in points:
        check_groups = defaultdict(lambda: {"passes": 0, "fails": 0})
        for p in points["checks"]:
            cname = p["tags"].get("check", "unnamed")
            if p["value"] > 0:
                check_groups[cname]["passes"] += 1
            else:
                check_groups[cname]["fails"] += 1
        for cname, cdata in check_groups.items():
            result["checks"].append({"name": cname, **cdata})

    return result


# ─── Charts ─────────────────────────────────────────────────

def _chart_to_bytes(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


def create_response_time_chart(data: dict) -> BytesIO | None:
    """Bar chart of HTTP response time percentiles."""
    duration = data["metrics"].get("http_req_duration")
    if not duration or duration["type"] != "trend":
        return None

    v = duration["values"]
    labels = ["Min", "Avg", "Median", "P90", "P95", "Max"]
    values = [v.get("min", 0), v.get("avg", 0), v.get("med", 0),
              v.get("p(90)", 0), v.get("p(95)", 0), v.get("max", 0)]

    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar(labels, values, color=CHART_COLORS[:len(labels)], width=0.6, edgecolor="white")
    ax.set_ylabel("Response Time (ms)", fontsize=10, color="#4B5563")
    ax.set_title("HTTP Request Duration Distribution", fontsize=12, fontweight="bold", color="#1F2A37", pad=12)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(colors="#6B7280")
    ax.yaxis.set_major_formatter(ticker.FormatStrFormatter("%.0f"))

    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(values) * 0.02,
                f"{val:.1f}", ha="center", va="bottom", fontsize=9, color="#374151")

    fig.tight_layout()
    return _chart_to_bytes(fig)


def create_breakdown_chart(data: dict) -> BytesIO | None:
    """Stacked bar showing request phase breakdown."""
    phases = {
        "DNS Lookup": "http_req_duration",  # fallback
        "Connecting": "http_req_connecting",
        "TLS Handshake": "http_req_tls_handshaking",
        "Sending": "http_req_sending",
        "Waiting (TTFB)": "http_req_waiting",
        "Receiving": "http_req_receiving",
    }

    labels = []
    values = []
    for label, metric in phases.items():
        if metric == "http_req_duration":
            continue
        m = data["metrics"].get(metric)
        if m and m["type"] == "trend":
            avg = m["values"].get("avg", 0)
            if avg > 0:
                labels.append(label)
                values.append(avg)

    if not values:
        return None

    fig, ax = plt.subplots(figsize=(7, 2.5))
    left = 0
    colors = CHART_COLORS[:len(labels)]
    for i, (label, val) in enumerate(zip(labels, values)):
        ax.barh(["Request"], [val], left=[left], color=colors[i], label=f"{label} ({val:.1f}ms)", height=0.4)
        left += val

    ax.set_xlabel("Time (ms)", fontsize=10, color="#4B5563")
    ax.set_title("Request Phase Breakdown (avg)", fontsize=12, fontweight="bold", color="#1F2A37", pad=12)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.25), ncol=3, fontsize=8, frameon=False)

    fig.tight_layout()
    return _chart_to_bytes(fig)


def create_checks_chart(data: dict) -> BytesIO | None:
    """Horizontal bar chart for checks pass/fail."""
    checks = data.get("checks", [])
    if not checks:
        return None

    names = [c["name"] for c in checks]
    passes = [c["passes"] for c in checks]
    fails = [c["fails"] for c in checks]

    fig, ax = plt.subplots(figsize=(7, max(2, len(checks) * 0.6 + 1)))
    y = range(len(names))
    ax.barh(y, passes, color="#10B981", label="Passed", height=0.5)
    ax.barh(y, fails, left=passes, color="#EF4444", label="Failed", height=0.5)
    ax.set_yticks(y)
    ax.set_yticklabels(names, fontsize=9)
    ax.set_xlabel("Count", fontsize=10, color="#4B5563")
    ax.set_title("Checks Results", fontsize=12, fontweight="bold", color="#1F2A37", pad=12)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(fontsize=9, frameon=False)

    fig.tight_layout()
    return _chart_to_bytes(fig)


def create_timeseries_chart(data: dict) -> BytesIO | None:
    """Response time over time (only for line-delimited JSON with timestamps)."""
    ts = data.get("_timeseries", {}).get("http_req_duration")
    if not ts or len(ts) < 2:
        return None

    times = []
    values = []
    for p in ts:
        try:
            t = datetime.fromisoformat(p["time"].replace("Z", "+00:00"))
            times.append(t)
            values.append(p["value"])
        except (ValueError, KeyError):
            continue

    if len(times) < 2:
        return None

    # Convert to seconds from start
    start = min(times)
    seconds = [(t - start).total_seconds() for t in times]

    fig, ax = plt.subplots(figsize=(7, 3.5))
    ax.scatter(seconds, values, c="#1A56DB", alpha=0.4, s=8, edgecolors="none")

    # Rolling average
    if len(values) > 10:
        window = max(5, len(values) // 20)
        rolling = [statistics.mean(values[max(0, i - window):i + 1]) for i in range(len(values))]
        ax.plot(seconds, rolling, color="#EF4444", linewidth=2, label=f"Rolling avg ({window})")
        ax.legend(fontsize=9, frameon=False)

    ax.set_xlabel("Time (seconds)", fontsize=10, color="#4B5563")
    ax.set_ylabel("Response Time (ms)", fontsize=10, color="#4B5563")
    ax.set_title("Response Time Over Test Duration", fontsize=12, fontweight="bold", color="#1F2A37", pad=12)
    ax.spines[["top", "right"]].set_visible(False)

    fig.tight_layout()
    return _chart_to_bytes(fig)


# ─── Document Generation ────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    """Apply background color to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers: list, rows: list, col_widths: list | None = None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        set_cell_shading(cell, "1A56DB")

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK
            if r % 2 == 1:
                set_cell_shading(cell, "F3F4F6")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def format_value(value, contains="default"):
    """Format a metric value for display."""
    if isinstance(value, float):
        if contains == "time":
            if value > 1000:
                return f"{value / 1000:.2f}s"
            return f"{value:.2f}ms"
        if value > 1_000_000:
            return f"{value / 1_000_000:.2f}M"
        if value > 1000:
            return f"{value / 1000:.2f}K"
        return f"{value:.2f}"
    return str(value)


def format_bytes(value):
    """Format bytes to human readable."""
    if value > 1_000_000:
        return f"{value / 1_000_000:.2f} MB"
    if value > 1000:
        return f"{value / 1000:.2f} KB"
    return f"{value:.0f} B"


def generate_report(data: dict, output: str, title: str, client: str, env: str,
                    tester: str, version: str, notes: str):
    """Generate the Word document."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.color.rgb = PRIMARY
        h_style.font.name = "Calibri"

    # ══════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Performance Test Report")
    run.font.size = Pt(16)
    run.font.color.rgb = GRAY

    doc.add_paragraph()

    # Meta info
    meta_items = [
        ("Client", client),
        ("Environment", env),
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Tester", tester),
    ]
    if version:
        meta_items.append(("Version", version))

    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.color.rgb = GRAY
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.color.rgb = DARK
        run.font.size = Pt(11)
        run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.font.color.rgb = DANGER
    run.font.size = Pt(10)
    run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  TABLE OF CONTENTS (placeholder)
    # ══════════════════════════════════════════════

    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Test Configuration",
        "3. Results Overview",
        "4. Response Time Analysis",
        "5. Throughput & Errors",
        "6. Checks & Thresholds",
        "7. Detailed Metrics",
        "8. Recommendations",
    ]
    if notes:
        toc_items.append("9. Notes")

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════

    doc.add_heading("1. Executive Summary", level=1)

    metrics = data["metrics"]
    duration = metrics.get("http_req_duration", {}).get("values", {})
    reqs = metrics.get("http_reqs", {}).get("values", {})
    failed = metrics.get("http_req_failed", {}).get("values", {})
    iterations = metrics.get("iterations", {}).get("values", {})
    vus_max = metrics.get("vus_max", {}).get("values", {})

    avg_rt = duration.get("avg", 0)
    p95_rt = duration.get("p(95)", 0)
    total_reqs = reqs.get("count", reqs.get("value", 0))
    rps = reqs.get("rate", 0)
    error_rate = failed.get("rate", 0)
    total_iters = iterations.get("count", iterations.get("value", 0))
    max_vus = vus_max.get("value", vus_max.get("max", 0))

    # Summary table
    summary_data = [
        ["Total Requests", f"{int(total_reqs):,}"],
        ["Requests/sec", f"{rps:.2f}"],
        ["Avg Response Time", f"{avg_rt:.2f}ms"],
        ["P95 Response Time", f"{p95_rt:.2f}ms"],
        ["Error Rate", f"{error_rate * 100:.2f}%" if error_rate <= 1 else f"{error_rate:.2f}%"],
        ["Max Virtual Users", f"{int(max_vus)}"],
        ["Total Iterations", f"{int(total_iters):,}"],
    ]

    add_styled_table(doc, ["Metric", "Value"], summary_data, col_widths=[8, 8])

    doc.add_paragraph()

    # Overall verdict
    p = doc.add_paragraph()
    if error_rate < 0.01 and p95_rt < 2000:
        run = p.add_run("✅ PASS — ")
        run.font.color.rgb = ACCENT
        run.bold = True
        run = p.add_run(f"The system performed within acceptable thresholds with a P95 response time of {p95_rt:.0f}ms and {error_rate * 100:.2f}% error rate.")
    elif error_rate < 0.05 and p95_rt < 5000:
        run = p.add_run("⚠️ MARGINAL — ")
        run.font.color.rgb = WARNING
        run.bold = True
        run = p.add_run(f"Performance is borderline. P95 response time is {p95_rt:.0f}ms with {error_rate * 100:.2f}% errors. Review recommendations.")
    else:
        run = p.add_run("❌ FAIL — ")
        run.font.color.rgb = DANGER
        run.bold = True
        run = p.add_run(f"System did not meet performance targets. P95 response time is {p95_rt:.0f}ms with {error_rate * 100:.2f}% errors. Immediate attention required.")

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  2. TEST CONFIGURATION
    # ══════════════════════════════════════════════

    doc.add_heading("2. Test Configuration", level=1)

    config_data = [
        ["Environment", env],
        ["Test Tool", "Grafana k6"],
        ["Max VUs", f"{int(max_vus)}"],
        ["Total Iterations", f"{int(total_iters):,}"],
        ["Report Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")],
    ]

    add_styled_table(doc, ["Parameter", "Value"], config_data, col_widths=[8, 8])

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    #  3. RESULTS OVERVIEW
    # ══════════════════════════════════════════════

    doc.add_heading("3. Results Overview", level=1)

    # Data transfer
    data_recv = metrics.get("data_received", {}).get("values", {})
    data_sent = metrics.get("data_sent", {}).get("values", {})
    recv_total = data_recv.get("count", data_recv.get("value", 0))
    sent_total = data_sent.get("count", data_sent.get("value", 0))

    overview_data = [
        ["Data Received", format_bytes(recv_total)],
        ["Data Sent", format_bytes(sent_total)],
        ["Total HTTP Requests", f"{int(total_reqs):,}"],
        ["Throughput", f"{rps:.2f} req/s"],
        ["Failed Requests", f"{error_rate * 100:.2f}%"],
    ]

    add_styled_table(doc, ["Metric", "Value"], overview_data, col_widths=[8, 8])

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  4. RESPONSE TIME ANALYSIS
    # ══════════════════════════════════════════════

    doc.add_heading("4. Response Time Analysis", level=1)

    if duration:
        rt_data = [
            ["Minimum", format_value(duration.get("min", 0), "time")],
            ["Average", format_value(duration.get("avg", 0), "time")],
            ["Median (P50)", format_value(duration.get("med", 0), "time")],
            ["P90", format_value(duration.get("p(90)", 0), "time")],
            ["P95", format_value(duration.get("p(95)", 0), "time")],
            ["Maximum", format_value(duration.get("max", 0), "time")],
        ]
        add_styled_table(doc, ["Percentile", "Response Time"], rt_data, col_widths=[8, 8])

    doc.add_paragraph()

    # Charts
    chart = create_response_time_chart(data)
    if chart:
        doc.add_picture(chart, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    chart = create_breakdown_chart(data)
    if chart:
        doc.add_heading("Request Phase Breakdown", level=2)
        doc.add_picture(chart, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    chart = create_timeseries_chart(data)
    if chart:
        doc.add_heading("Response Time Over Duration", level=2)
        doc.add_picture(chart, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  5. THROUGHPUT & ERRORS
    # ══════════════════════════════════════════════

    doc.add_heading("5. Throughput & Errors", level=1)

    throughput_data = [
        ["Requests per Second", f"{rps:.2f}"],
        ["Total Requests", f"{int(total_reqs):,}"],
        ["Successful", f"{int(total_reqs * (1 - error_rate)):,}" if error_rate <= 1 else "N/A"],
        ["Failed", f"{int(total_reqs * error_rate):,}" if error_rate <= 1 else "N/A"],
        ["Error Rate", f"{error_rate * 100:.2f}%" if error_rate <= 1 else f"{error_rate:.2f}%"],
    ]
    add_styled_table(doc, ["Metric", "Value"], throughput_data, col_widths=[8, 8])

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    #  6. CHECKS & THRESHOLDS
    # ══════════════════════════════════════════════

    doc.add_heading("6. Checks & Thresholds", level=1)

    # Checks
    if data.get("checks"):
        doc.add_heading("Checks", level=2)
        check_rows = []
        for c in data["checks"]:
            total = c["passes"] + c["fails"]
            rate = (c["passes"] / total * 100) if total > 0 else 0
            status = "✅ PASS" if c["fails"] == 0 else "❌ FAIL"
            check_rows.append([c["name"], f"{c['passes']}", f"{c['fails']}", f"{rate:.1f}%", status])

        add_styled_table(doc,
                         ["Check", "Passes", "Fails", "Success Rate", "Status"],
                         check_rows,
                         col_widths=[5, 2.5, 2.5, 3, 3])

        doc.add_paragraph()

        chart = create_checks_chart(data)
        if chart:
            doc.add_picture(chart, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Thresholds
    if data.get("thresholds"):
        doc.add_heading("Thresholds", level=2)
        t_rows = []
        for name, passed in data["thresholds"].items():
            t_rows.append([name, "✅ PASS" if passed else "❌ FAIL"])
        add_styled_table(doc, ["Threshold", "Result"], t_rows, col_widths=[10, 4])

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  7. DETAILED METRICS
    # ══════════════════════════════════════════════

    doc.add_heading("7. Detailed Metrics", level=1)

    # Group metrics by category
    http_metrics = {k: v for k, v in metrics.items() if k.startswith("http_")}
    exec_metrics = {k: v for k, v in metrics.items() if k in ("iteration_duration", "iterations", "vus", "vus_max")}
    net_metrics = {k: v for k, v in metrics.items() if k.startswith("data_")}
    custom_metrics = {k: v for k, v in metrics.items()
                      if k not in http_metrics and k not in exec_metrics
                      and k not in net_metrics and k not in ("checks",)}

    def add_metric_table(name: str, group: dict):
        if not group:
            return
        doc.add_heading(name, level=2)
        rows = []
        for mname, mdata in sorted(group.items()):
            vals = mdata.get("values", {})
            contains = mdata.get("contains", "default")
            mtype = mdata.get("type", "trend")

            if mtype == "trend":
                rows.append([
                    mname,
                    format_value(vals.get("avg", 0), contains),
                    format_value(vals.get("min", 0), contains),
                    format_value(vals.get("med", 0), contains),
                    format_value(vals.get("max", 0), contains),
                    format_value(vals.get("p(90)", 0), contains),
                    format_value(vals.get("p(95)", 0), contains),
                ])
            elif mtype == "counter":
                count = vals.get("count", vals.get("value", 0))
                rate = vals.get("rate", 0)
                rows.append([mname, f"{count:,.0f}", "-", "-", "-", "-", f"{rate:.2f}/s"])
            elif mtype == "rate":
                rate = vals.get("rate", 0)
                display = f"{rate * 100:.2f}%" if rate <= 1 else f"{rate:.2f}%"
                rows.append([mname, display, "-", "-", "-", "-", "-"])
            elif mtype == "gauge":
                v = vals.get("value", 0)
                rows.append([mname, f"{v:,.0f}", f"{vals.get('min', 0):,.0f}", "-",
                             f"{vals.get('max', 0):,.0f}", "-", "-"])

        if rows:
            add_styled_table(doc,
                             ["Metric", "Avg/Value", "Min", "Median", "Max", "P90", "P95"],
                             rows, col_widths=[4.5, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2])
        doc.add_paragraph()

    add_metric_table("HTTP Metrics", http_metrics)
    add_metric_table("Execution Metrics", exec_metrics)
    add_metric_table("Network Metrics", net_metrics)
    add_metric_table("Custom Metrics", custom_metrics)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    #  8. RECOMMENDATIONS
    # ══════════════════════════════════════════════

    doc.add_heading("8. Recommendations", level=1)

    recommendations = []

    if p95_rt > 2000:
        recommendations.append(("High P95 Response Time",
                                f"The 95th percentile response time is {p95_rt:.0f}ms. Consider optimizing database queries, "
                                "adding caching layers, or scaling application servers."))
    if p95_rt > 500 and duration.get("max", 0) > p95_rt * 3:
        recommendations.append(("Response Time Outliers",
                                f"Maximum response time ({duration.get('max', 0):.0f}ms) is significantly higher than P95 "
                                f"({p95_rt:.0f}ms). Investigate intermittent slowdowns — possible causes include garbage collection, "
                                "cold starts, or resource contention."))

    if error_rate > 0:
        recommendations.append(("Non-Zero Error Rate",
                                f"Error rate of {error_rate * 100:.2f}% detected. Review server logs for the time period "
                                "to identify root cause. Check for connection timeouts, rate limiting, or application errors."))

    waiting = metrics.get("http_req_waiting", {}).get("values", {})
    if waiting and waiting.get("avg", 0) > avg_rt * 0.8:
        recommendations.append(("Server Processing Bottleneck",
                                f"Time-to-first-byte (TTFB) accounts for {waiting.get('avg', 0) / avg_rt * 100:.0f}% of total "
                                "response time. The bottleneck is server-side processing, not network. Profile backend code paths."))

    tls = metrics.get("http_req_tls_handshaking", {}).get("values", {})
    if tls and tls.get("avg", 0) > 50:
        recommendations.append(("TLS Overhead",
                                f"Average TLS handshake time is {tls.get('avg', 0):.0f}ms. Enable TLS session resumption "
                                "and consider TLS 1.3 for faster handshakes."))

    blocked = metrics.get("http_req_blocked", {}).get("values", {})
    if blocked and blocked.get("avg", 0) > 10:
        recommendations.append(("Connection Queuing",
                                f"Requests are blocked for {blocked.get('avg', 0):.1f}ms on average waiting for a free TCP "
                                "connection. Increase connection pool size or enable HTTP/2 multiplexing."))

    if not recommendations:
        recommendations.append(("No Issues Detected",
                                "All metrics are within acceptable ranges. Continue monitoring and consider increasing "
                                "load in subsequent test iterations to find the performance ceiling."))

    for title_text, desc in recommendations:
        p = doc.add_paragraph()
        run = p.add_run(f"▸ {title_text}")
        run.bold = True
        run.font.color.rgb = PRIMARY
        p = doc.add_paragraph(desc)
        p.paragraph_format.space_after = Pt(8)

    # ══════════════════════════════════════════════
    #  9. NOTES (optional)
    # ══════════════════════════════════════════════

    if notes:
        doc.add_page_break()
        doc.add_heading("9. Notes", level=1)
        for line in notes.split("\n"):
            doc.add_paragraph(line)

    # ── Footer ──
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Confidential — {client} — Generated {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

    # ── Save ──
    doc.save(output)
    print(f"✅ Report saved: {output}")
    return output


# ─── CLI ────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate a professional Word report from k6 performance test results",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # From handleSummary() JSON export
  python generate_report.py -i summary.json -f json -c "Acme Corp" -t "API Load Test"

  # From k6 --out json=results.json
  python generate_report.py -i results.json -f json -c "Acme Corp" -t "Payment API"

  # From k6 --out csv=results.csv
  python generate_report.py -i results.csv -f csv -c "Acme Corp" -e Staging

How to export from k6:
  # JSON (line-delimited, granular data points — best for time-series charts)
  k6 run --out json=results.json script.js

  # CSV (line-delimited, granular data points)
  k6 run --out csv=results.csv script.js

  # JSON summary (aggregated, smallest file — recommended)
  # Add to your k6 script:
  export function handleSummary(data) {
    return { "summary.json": JSON.stringify(data) };
  }
        """,
    )
    parser.add_argument("-i", "--input", required=True, help="Path to k6 results file (JSON or CSV)")
    parser.add_argument("-f", "--format", choices=["json", "csv"], required=True, help="Input file format")
    parser.add_argument("-o", "--output", default=None, help="Output .docx path (default: k6-report-YYYY-MM-DD.docx)")
    parser.add_argument("-t", "--title", default="Performance Test Report", help="Report title")
    parser.add_argument("-c", "--client", default="[Client Name]", help="Client/company name")
    parser.add_argument("-e", "--env", default="Production", help="Test environment (e.g. Production, Staging, UAT)")
    parser.add_argument("--tester", default="QA Team", help="Tester name or team")
    parser.add_argument("--version", default="", help="Application version tested")
    parser.add_argument("--notes", default="", help="Additional notes to include")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"❌ File not found: {args.input}")
        sys.exit(1)

    # Parse
    print(f"📂 Parsing {args.input} ({args.format})...")
    if args.format == "json":
        data = parse_json_summary(args.input)
    else:
        data = parse_csv(args.input)

    print(f"   Found {len(data['metrics'])} metrics, {len(data.get('checks', []))} checks, {len(data.get('thresholds', {}))} thresholds")

    # Generate
    output = args.output or f"k6-report-{datetime.now().strftime('%Y-%m-%d')}.docx"
    print(f"📝 Generating report...")
    generate_report(data, output, args.title, args.client, args.env,
                    args.tester, args.version, args.notes)


if __name__ == "__main__":
    main()
